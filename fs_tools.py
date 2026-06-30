"""
fs_tools.py
===========
Part A — Core File System Tools (Milestone 1)

Provides 4 LLM-callable tools for file I/O operations on resume documents:
  - read_file(filepath)            → read PDF / TXT / DOCX content
  - list_files(directory, ext)     → list files with metadata
  - write_file(filepath, content)  → write content to disk
  - search_in_file(filepath, kw)   → keyword search with context

Each tool returns a structured dict (never raises), so an LLM orchestrator
can safely call them and branch on `success`/`error` without try/except.
"""

import os
import re
import json
from datetime import datetime
from typing import Optional, List, Dict, Any

# ── Optional parsers (graceful degradation if not installed) ─────────────────
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx as docx_lib
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Tool 1: read_file ──────────────────────────────────────────────────────

def read_file(filepath: str) -> dict:
    """
    Read a resume file (PDF, TXT, or DOCX) and extract its text content.

    Args:
        filepath: Path to the file to read.

    Returns:
        {
            "success": bool,
            "content": str,           # extracted text (empty on failure)
            "metadata": {
                "filename": str,
                "extension": str,
                "size_bytes": int,
                "char_count": int,
                "word_count": int,
            },
            "error": str | None
        }
    """
    if not os.path.exists(filepath):
        return _error_response(f"File not found: {filepath}")

    if not os.path.isfile(filepath):
        return _error_response(f"Path is not a file: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)
    size_bytes = os.path.getsize(filepath)

    try:
        if ext == ".pdf":
            content = _read_pdf(filepath)
        elif ext == ".docx":
            content = _read_docx(filepath)
        elif ext in (".txt", ".md"):
            content = _read_text(filepath)
        else:
            return _error_response(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .md")

        return {
            "success": True,
            "content": content,
            "metadata": {
                "filename": filename,
                "extension": ext,
                "size_bytes": size_bytes,
                "char_count": len(content),
                "word_count": len(content.split()),
            },
            "error": None
        }

    except Exception as e:
        return _error_response(f"Failed to read {filename}: {str(e)}")


def _read_text(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(filepath: str) -> str:
    if not PDF_AVAILABLE:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")
    reader = PdfReader(filepath)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _read_docx(filepath: str) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = docx_lib.Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    # Also pull table content (resumes sometimes use tables for skills)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


# ── Tool 2: list_files ─────────────────────────────────────────────────────

def list_files(directory: str, extension: Optional[str] = None) -> list:
    """
    List all files in a directory, optionally filtered by extension.

    Args:
        directory: Directory path to scan.
        extension: Optional filter, e.g. ".pdf", ".txt" (case-insensitive,
                   leading dot optional).

    Returns:
        List of dicts:
        [
            {
                "filename": str,
                "filepath": str,
                "extension": str,
                "size_bytes": int,
                "modified_date": str (ISO format),
            },
            ...
        ]
        Returns [] if directory doesn't exist or has no matching files.
    """
    if not os.path.exists(directory) or not os.path.isdir(directory):
        return []

    if extension and not extension.startswith("."):
        extension = "." + extension
    if extension:
        extension = extension.lower()

    results = []
    for fname in sorted(os.listdir(directory)):
        fpath = os.path.join(directory, fname)
        if not os.path.isfile(fpath):
            continue

        ext = os.path.splitext(fname)[1].lower()
        if extension and ext != extension:
            continue

        stat = os.stat(fpath)
        results.append({
            "filename": fname,
            "filepath": fpath,
            "extension": ext,
            "size_bytes": stat.st_size,
            "modified_date": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        })

    return results


# ── Tool 3: write_file ─────────────────────────────────────────────────────

def write_file(filepath: str, content: str) -> dict:
    """
    Write content to a file, creating parent directories as needed.

    Args:
        filepath: Destination path.
        content: Text content to write.

    Returns:
        {
            "success": bool,
            "filepath": str,
            "bytes_written": int,
            "error": str | None
        }
    """
    try:
        parent_dir = os.path.dirname(filepath)
        if parent_dir and not os.path.exists(parent_dir):
            os.makedirs(parent_dir, exist_ok=True)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return {
            "success": True,
            "filepath": filepath,
            "bytes_written": len(content.encode("utf-8")),
            "error": None
        }

    except Exception as e:
        return {
            "success": False,
            "filepath": filepath,
            "bytes_written": 0,
            "error": str(e)
        }


# ── Tool 4: search_in_file ─────────────────────────────────────────────────

def search_in_file(filepath: str, keyword: str, context_chars: int = 80) -> dict:
    """
    Search for a keyword inside a file's content (case-insensitive).

    Args:
        filepath: File to search in (PDF, TXT, DOCX — reused via read_file).
        keyword: Term to search for.
        context_chars: Number of characters of surrounding context to include
                        on each side of a match.

    Returns:
        {
            "success": bool,
            "filepath": str,
            "keyword": str,
            "match_count": int,
            "matches": [
                {"position": int, "context": str},
                ...
            ],
            "error": str | None
        }
    """
    read_result = read_file(filepath)
    if not read_result["success"]:
        return {
            "success": False,
            "filepath": filepath,
            "keyword": keyword,
            "match_count": 0,
            "matches": [],
            "error": read_result["error"]
        }

    content = read_result["content"]
    pattern = re.escape(keyword)
    matches = []

    for m in re.finditer(pattern, content, re.IGNORECASE):
        start = max(0, m.start() - context_chars)
        end = min(len(content), m.end() + context_chars)
        snippet = content[start:end].replace("\n", " ").strip()
        prefix = "..." if start > 0 else ""
        suffix = "..." if end < len(content) else ""
        matches.append({
            "position": m.start(),
            "context": f"{prefix}{snippet}{suffix}"
        })

    return {
        "success": True,
        "filepath": filepath,
        "keyword": keyword,
        "match_count": len(matches),
        "matches": matches,
        "error": None
    }


# ── Helpers ────────────────────────────────────────────────────────────────

def _error_response(message: str) -> dict:
    return {
        "success": False,
        "content": "",
        "metadata": {},
        "error": message
    }


# ── Tool Schema (for LLM function calling) ───────────────────────────────────

TOOL_SCHEMAS = [
    {
        "name": "read_file",
        "description": "Read a resume file (PDF, DOCX, or TXT) and extract its text content along with metadata.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path to the file to read."}
            },
            "required": ["filepath"]
        }
    },
    {
        "name": "list_files",
        "description": "List all files in a directory, optionally filtered by file extension.",
        "input_schema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Directory path to scan."},
                "extension": {"type": "string", "description": "Optional extension filter, e.g. '.pdf' or '.txt'."}
            },
            "required": ["directory"]
        }
    },
    {
        "name": "write_file",
        "description": "Write text content to a file, creating any needed parent directories.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Destination file path."},
                "content": {"type": "string", "description": "Text content to write."}
            },
            "required": ["filepath", "content"]
        }
    },
    {
        "name": "search_in_file",
        "description": "Search for a keyword inside a file's content (case-insensitive) and return matches with surrounding context.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "File to search in."},
                "keyword": {"type": "string", "description": "Term to search for."}
            },
            "required": ["filepath", "keyword"]
        }
    }
]


# Map tool name → function, for dynamic dispatch by the LLM assistant
TOOL_REGISTRY = {
    "read_file": read_file,
    "list_files": list_files,
    "write_file": write_file,
    "search_in_file": search_in_file,
}


# ── Self-test ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("  fs_tools.py — Self Test")
    print("=" * 60)

    # Test write_file
    print("\n1. write_file()")
    result = write_file("resumes/_test.txt", "Test content with Python and AWS skills.")
    print(f"   {result}")

    # Test list_files
    print("\n2. list_files()")
    files = list_files("resumes", extension=".txt")
    print(f"   Found {len(files)} .txt files")
    for f in files[:3]:
        print(f"   - {f['filename']} ({f['size_bytes']} bytes)")

    # Test read_file
    print("\n3. read_file()")
    result = read_file("resumes/_test.txt")
    print(f"   success={result['success']}, chars={result['metadata'].get('char_count')}")

    # Test search_in_file
    print("\n4. search_in_file()")
    result = search_in_file("resumes/_test.txt", "Python")
    print(f"   match_count={result['match_count']}")
    for m in result["matches"]:
        print(f"   - {m['context']}")

    # Cleanup test file
    os.remove("resumes/_test.txt")
    print("\n✅ All self-tests completed")
